import warnings

from docx import Document
from docx.text.paragraph import Paragraph
from tqdm import tqdm
import glob
import spacy
import pandas as pd
from name_matching.name_matcher import NameMatcher
import os
from utils import parse_docx, sanitize_str, split_text_into_chunks
from docx.table import Table
import copy
from pathlib import Path
warnings.simplefilter(action='ignore', category=FutureWarning)
class DataRedaction:
    def __init__(self,input_dir_path,redacted_output_path):
        self.entity_types_to_mask = ["DATE", "MONEY","PERCENT", "QUANTITY", "TIME"]
        self.entity_type_to_replace = ["ORG","PERSON","PRODUCT"]
        self.nlp = spacy.load("en_core_web_trf")
        self.input_dir_path = input_dir_path
        self.redacted_output_path = redacted_output_path
        os.makedirs(self.redacted_output_path,exist_ok=True)
        self.number_of_matches = 10

        self.matcher = NameMatcher(top_n=self.number_of_matches * 2,
                                number_of_matches = self.number_of_matches,
                              lowercase=True,
                              punctuations=True,
                              remove_ascii=True,
                              legal_suffixes=True,
                              common_words=False,
                              verbose=False)
        self.matcher.set_distance_metrics(['discounted_levenshtein',
                                      'SSK',
                                      'fuzzy_wuzzy_token_sort'])
        self.matching_threshold = 80 # match below this would be discarded
        self.table_insertion_pattern = '\n\n####INSERT TABLE HERE####\n\n'

    def _redact_entity(self, ent,grouped_entities:dict,last_group_id) -> tuple[str,int]:
        if ent.label_ in self.entity_type_to_replace:
            ent_text_type = ent.text.lower().strip() + '_' + ent.label_
            group_id = grouped_entities.get(ent_text_type)
            if group_id is not None:
                redacted_text = ent.label_ + "_" + str(group_id)
            else:
                redacted_text = ent.label_ + "_" + str(last_group_id)
                last_group_id += 1
        else:
            redacted_text = "XXXXX"

        return redacted_text,last_group_id

    def _group_similar_names(self, entities:list) -> dict:
        # Groups the entities based on similarity and returns a dict with entity text as key and group id as value
        group_id = 0
        grouped_entities = {}
        for i in range(len(entities)):
            entity_text = entities[i].text.lower().strip()
            entity_text_type = entity_text + '_' + entities[i].label_

            if grouped_entities.get(entity_text_type) or (entities[i].label_ not in self.entity_type_to_replace and entities[i].label_ not in self.entity_types_to_mask):
                continue

            similar_type_ents = [j for j in entities[i+1:] if j.label_ == entities[i].label_ and
                                 entity_text != j.text.lower().strip() and
                                 j.label_ in self.entity_type_to_replace]
            if len(similar_type_ents) > 0:
                names_to_be_matched = pd.DataFrame({"name":[n.text for n in similar_type_ents],
                                                    "ent":similar_type_ents})
                names_to_be_matched.drop_duplicates(subset=['name'],inplace=True)
                name_to_match = pd.DataFrame({"name":[entities[i].text],
                                              "ent":[entities[i]]})
                self.matcher.load_and_process_master_data('name', names_to_be_matched)
                result = self.matcher.match_names(to_be_matched=name_to_match, column_matching='name').iloc[0]

                grouped_entities[entity_text_type] = group_id

                for j in range(self.number_of_matches):
                    if result['score_'+str(j)] >= self.matching_threshold:
                        match_entity_text_type = result['match_name_'+str(j)].lower().strip() + '_' + entities[i].label_
                        grouped_entities[match_entity_text_type] = group_id
                group_id += 1

        return grouped_entities

    def _redact_extracted_entities(self, doc, entities,tables) -> tuple[Document,pd.DataFrame]:
        redacted_text_map = {}
        try:
            entity_groups = self._group_similar_names(entities)
        except:
            entity_groups = {}

        redacted_doc = Document()
        paragraph = redacted_doc.add_paragraph('')

        redacted_text = doc.text[0:entities[0].start_char]

        paragraph = self.insert_table_if_needed(redacted_text, redacted_doc,paragraph, tables)

        original_redacted_map = []
        if len(entity_groups.values()) > 0 :
            last_group_id = max(entity_groups.values()) + 1
        else:
            last_group_id = 0
        for i, ent in enumerate(entities):
            if ent.label_ in self.entity_types_to_mask or ent.label_ in self.entity_type_to_replace:

                text_based_replacement = redacted_text_map.get(ent.text.lower().strip()+ent.label_)
                if text_based_replacement is not None:
                    redacted_entity_text = text_based_replacement
                else:
                    redacted_entity_text, last_group_id = self._redact_entity(ent, entity_groups,last_group_id)
                    redacted_text_map[ent.text.lower().strip()+ent.label_] = redacted_entity_text
                paragraph.add_run(redacted_entity_text).add_comment(ent.text)

                #redacted_text = redacted_text + redacted_entity_text + " (((" + ent.text+"))) "
                if i < len(entities) - 1:
                    text_till_next_entity = doc.text[ent.end_char:entities[i + 1].start_char]

                else:
                    text_till_next_entity = doc.text[entities[-1].end_char:]

                paragraph = self.insert_table_if_needed(text_till_next_entity, redacted_doc,paragraph, tables)
                original_redacted_map.append({"original_text": ent.text,
                                              "original_span": (ent.start_char, ent.end_char),
                                              "entity_type": ent.label_,
                                              "redacted_text": redacted_entity_text})

            else:
                # Add the entity text as it is
                if i < len(entities) - 1:
                    text_till_next_entity = doc.text[ent.end_char:entities[i + 1].start_char]

                else:
                    text_till_next_entity = doc.text[entities[-1].end_char:]

                paragraph = self.insert_table_if_needed(text_till_next_entity, redacted_doc,paragraph, tables)

        original_redacted_df = pd.DataFrame(original_redacted_map)
        return redacted_doc, original_redacted_df

    def _get_and_create_relative_output_dir(self, input_path):
        # Creates output directories in out put folder relative to input path
        relative_path = Path(input_path).parent.relative_to(self.input_dir_path)
        output_docs_file_path = Path(self.redacted_output_path).joinpath(relative_path)
        output_docs_file_path.mkdir(parents=True, exist_ok=True)
        return str(output_docs_file_path)
    def _insert_tables_back(self, redacted_doc, tables) -> Document:
        table_inserted_redacted_doc = Document()
        text_chunks = redacted_text.split(self.table_insertion_pattern)
        if len(text_chunks) > 1:
            for i in range(len(text_chunks)):
                para = table_inserted_redacted_doc.add_paragraph(sanitize_str(text_chunks[i]))
                if i < len(tables):
                    new_table = copy.deepcopy(tables[i]._tbl)
                    para._p.addnext(new_table)
        else:
            table_inserted_redacted_doc.add_paragraph(sanitize_str(text_chunks[0]))

        return table_inserted_redacted_doc


    def redact_all_files_in_folder(self):
        for file_path in tqdm( glob.glob(self.input_dir_path + '/**/*.docx',recursive=True),desc = 'Data Redaction'):
            self.redact_one_file(file_path)

    def redact_one_file(self,file_path):
        # try:
            parsed_docx = parse_docx(file_path)
            text_chunks = [i for i in parsed_docx if type(i) == str]
            tables = [i for i in parsed_docx if type(i) == Table]
            tables.reverse()
            combined_text = self.table_insertion_pattern.join(text_chunks)
            text_chunks = split_text_into_chunks(combined_text)
            doc_list = []
            for chunk in text_chunks:
                doc_list.append(self.nlp(chunk))
            doc = spacy.tokens.Doc.from_docs(doc_list)
            entities = [i for i in doc.ents]
            entities = sorted(entities, key=lambda x: x.start_char)

            redacted_doc,original_redacted_df = self._redact_extracted_entities(doc, entities,tables)

            #redacted_doc = self._insert_tables_back(redacted_doc, tables)

            output_dir = self._get_and_create_relative_output_dir(file_path)
            output_doc_path = os.path.join(output_dir, os.path.basename(file_path).replace('.docx', '_redacted.docx'))
            redacted_doc.save(output_doc_path)
            redaction_map_file_name = os.path.basename(file_path).replace('.docx', '_redaction_map.csv')
            original_redacted_df.to_csv(os.path.join(output_dir, redaction_map_file_name), index=False)
        # except:
        #     print("Could not process " + file_path)

    def insert_table_if_needed(self, redacted_text, redacted_doc, paragraph, tables):
        if redacted_text.__contains__(self.table_insertion_pattern):
            text_chunks = redacted_text.split(self.table_insertion_pattern)
            for i in range(len(text_chunks)-1):
                #paragraph = redacted_doc.add_paragraph(sanitize_str(text_chunks[i]))
                paragraph.add_run(sanitize_str(text_chunks[i]))

                if len(tables)>0:
                    new_table = copy.deepcopy(tables.pop()._tbl)
                    paragraph_new = redacted_doc.add_paragraph('')
                    paragraph_new._p.addnext(new_table)
                    paragraph = redacted_doc.add_paragraph('')

        else:
            paragraph.add_run(sanitize_str(redacted_text))
        return paragraph

