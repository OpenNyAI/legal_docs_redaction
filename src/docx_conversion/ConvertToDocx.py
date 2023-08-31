from docx import Document
from matplotlib import pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
import deepdoctection as dd
import glob
import os
import pdfplumber
import pdf2image
import pytesseract
from joblib import Parallel, delayed
from tqdm import tqdm
from src.utils import parse_pdf
from pathlib import Path
class ConvertToDocx:
    def __init__(self, input_folder, output_docx_folder_path):
        self.input_folder = input_folder
        self.output_docx_folder_path = output_docx_folder_path

        os.makedirs(output_docx_folder_path, exist_ok=True)

        self.use_layout_detection = False # If false then plain tesseract would be used. Else layout detection is done
        self.model_for_layout_detection = 'detectron2' # 'table-transformer' or 'detectron2'

        if self.use_layout_detection:
            self.annotated_pdf_folder_path = os.path.join(output_docx_folder_path, 'annotated')
            os.makedirs(self.annotated_pdf_folder_path, exist_ok=True)

    def _build_deepdoctection_analyzer(self):
        config_overwrite = ["USE_LAYOUT=False",
                            "USE_TABLE_SEGMENTATION=False",
                            "USE_TABLE_REFINEMENT=False"]

        if self.use_layout_detection:
            if self.model_for_layout_detection == 'table-transformer':
                # Run pip install deepdoctection[pt]
                config_overwrite = ["PT.LAYOUT.WEIGHTS=microsoft/table-transformer-detection/pytorch_model.bin",
                                    "USE_TABLE_SEGMENTATION=True",
                                    "PT.LAYOUT.PAD.TOP=1",
                                    "PT.LAYOUT.PAD.RIGHT=1",
                                    "PT.LAYOUT.PAD.BOTTOM=1",
                                    "PT.LAYOUT.PAD.LEFT=1"]

            elif self.model_for_layout_detection == 'detectron2':
                config_overwrite = ["USE_LAYOUT=True",
                                    "USE_TABLE_SEGMENTATION=True",
                                    "USE_TABLE_REFINEMENT=True",
                                    "TEXT_ORDERING.INCLUDE_RESIDUAL_TEXT_CONTAINER=True"]

        analyzer = dd.get_dd_analyzer(config_overwrite=config_overwrite)
        return analyzer
    def _convert_pdf_to_text_ocr_with_layout_detection(self,analyzer, input_pdf_path,annotated_pdf_path)->str:
        df = analyzer.analyze(path=input_pdf_path)
        df.reset_state()

        doc = iter(df)

        doc_text = ''
        pp = PdfPages(annotated_pdf_path)

        for i,page in enumerate(doc):
            fig = plt.figure(figsize=(25,17))
            plt.axis('off')
            plt.imshow(page.viz())

            pp.savefig(fig)
            plt.close(fig)
            doc_text += page.text
            if i > 10:
                break
        pp.close()
        return doc_text

    def _extract_text_without_ocr(self,input_pdf_path) -> str:
        with pdfplumber.open(input_pdf_path) as pdf:
            doc_text = ''
            for page in pdf.pages:
                doc_text += page.extract_text()
        return doc_text


    def _get_images_from_pdf(self,file_path: str, dpi: int = 200, fmt: str = 'png', thread_count: int = 1) -> list:
        return pdf2image.convert_from_path(file_path, dpi=dpi, fmt=fmt, thread_count=thread_count)
    def _extract_text_from_image(self,image_bytes, custom_oem_psm_config=r'--psm 6 -c preserve_interword_spaces=1',
                                lang="eng"):
        return pytesseract.image_to_string(image_bytes, lang=lang, config=custom_oem_psm_config)

    def _convert_pdf_to_text_ocr(self,pdf_path,lang = 'eng') -> str:
        if self.use_layout_detection:
            analyzer = self._build_deepdoctection_analyzer()
            annotated_pdf_path = os.path.join(self.annotated_pdf_folder_path,os.path.basename(pdf_path))
            pdf_text = self._convert_pdf_to_text_ocr_with_layout_detection(analyzer,pdf_path,annotated_pdf_path)
        else:
            # Use plain Tesseract
            pdf_text = ''
            pdf_images = self._get_images_from_pdf(pdf_path)
            for image in pdf_images:
                text = self._extract_text_from_image(image,lang=lang)
                pdf_text += text
        return pdf_text

    def _convert_pdf_to_docx_and_write(self, file_path):
        # First check if the pdf is image based or text based
        try:
            parsed_pdf = parse_pdf(file_path)
            if len(parsed_pdf) == 0:
                # pdf is image based, use OCR
                doc_text = self._convert_pdf_to_text_ocr(file_path)
                document = Document()
                document.add_paragraph(doc_text)
            else:
                document = Document()
                for element in parsed_pdf:
                    if type(element) == str:
                        document.add_paragraph(element)
                    elif type(element) == list:
                        row_cnt = len(element)
                        col_cnt = len(element[0])
                        table = document.add_table(rows = row_cnt, cols = col_cnt)
                        table.style = 'Table Grid'
                        for row in range(row_cnt):
                            for col in range(col_cnt):
                                try:
                                    cell_text = element[row][col]
                                except:
                                    cell_text  = ''
                                table.cell(row, col).text = cell_text

            output_dir = self._get_and_create_relative_output_dir(file_path)
            output_docs_file_path = os.path.join(output_dir,os.path.basename(file_path).replace('.pdf','.docx'))
            document.save(output_docs_file_path)
        except:
            print("Could not convert PDF to DOCX. Skipping  " + file_path)

    def _get_and_create_relative_output_dir(self, input_path):
        # Creates output directories in out put folder relative to input path
        relative_path = Path(input_path).parent.relative_to(self.input_folder)
        output_docs_file_path = Path(self.output_docx_folder_path).joinpath(relative_path)
        output_docs_file_path.mkdir(parents=True, exist_ok=True)
        return str(output_docs_file_path)
    def _convert_doc_to_docx_and_write(self, doc_file_path):
        try:
            output_dir = self._get_and_create_relative_output_dir(doc_file_path)
            os.system(f'soffice --headless --convert-to docx "{doc_file_path}" --outdir "{output_dir}"')
        except:
            print("Could not convert DOC to DOCX. Skipping  " + doc_file_path)
    def convert_to_docx(self):
        # 1. Convert PDF files to docx
        Parallel(n_jobs=-1)(
            delayed(self._convert_pdf_to_docx_and_write)(path) for path in
            tqdm(glob.glob(self.input_folder + '/**/*.pdf',recursive=True),desc = 'Converting PDF to DOCX'))

        # 2. Convert doc files to docx
        Parallel(n_jobs=-1)(
            delayed(self._convert_doc_to_docx_and_write)(path) for path in
            tqdm(glob.glob(self.input_folder + '/**/*.doc',recursive=True),desc = 'Converting DOC to DOCX'))

        # 3. copy docx files to output folder
        for file_path in tqdm(glob.glob(self.input_folder + '/**/*.docx',recursive=True)):
            output_dir = self._get_and_create_relative_output_dir(file_path)
            os.system(f'cp "{file_path}" "{output_dir}"')

if __name__ == '__main__':
    input_pdfs_folder = '/Users/prathamesh/tw_projects/OpenNyAI/data/LLM/data_readaction/test'
    output_text_folder_path = input_pdfs_folder+'_docx'

    docx_converter = ConvertToDocx(input_pdfs_folder, output_text_folder_path)
    docx_converter.convert_to_docx()



