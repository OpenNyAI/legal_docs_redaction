from docx import Document
import glob
import os
import pdfplumber
import pdf2image
import pytesseract
from joblib import Parallel, delayed
from tqdm import tqdm
from utils import parse_pdf, sanitize_str
from pathlib import Path
import re
class ConvertToDocx:
    def __init__(self, input_folder, output_docx_folder_path):
        self.input_folder = input_folder
        self.output_docx_folder_path = output_docx_folder_path

        os.makedirs(output_docx_folder_path, exist_ok=True)


    def _extract_text_without_ocr(self,input_pdf_path) -> str:
        with pdfplumber.open(input_pdf_path) as pdf:
            doc_text = ''
            for page in pdf.pages:
                doc_text += page.extract_text()
        return doc_text


    def _get_images_from_pdf(self,file_path: str, dpi: int = 200, fmt: str = 'png', thread_count: int = 1) -> list:
        return pdf2image.convert_from_path(file_path, dpi=dpi, fmt=fmt, thread_count=thread_count)
    def _extract_text_from_image(self,image_bytes, custom_oem_psm_config=r'--psm 6 -c preserve_interword_spaces=1',
                                lang="eng"):
        return pytesseract.image_to_string(image_bytes, lang=lang, config=custom_oem_psm_config)

    def _convert_pdf_to_text_ocr(self,pdf_path,lang = 'eng') -> str:
        # Use plain Tesseract
        pdf_text = ''
        pdf_images = self._get_images_from_pdf(pdf_path)
        for image in pdf_images:
            text = self._extract_text_from_image(image,lang=lang)
            pdf_text += text
        return pdf_text

    def _convert_pdf_to_docx_and_write(self, file_path):
        # First check if the pdf is image based or text based
        try:
            parsed_pdf = parse_pdf(file_path)
            if len(parsed_pdf) == 0:
                # pdf is image based, use OCR
                doc_text = self._convert_pdf_to_text_ocr(file_path)
                document = Document()
                document.add_paragraph(sanitize_str(doc_text))
            else:
                document = Document()
                for element in parsed_pdf:
                    if type(element) == str:
                        document.add_paragraph(sanitize_str(element))
                    elif type(element) == list:
                        row_cnt = len(element)
                        col_cnt = len(element[0])
                        table = document.add_table(rows = row_cnt, cols = col_cnt)
                        table.style = 'Table Grid'
                        for row in range(row_cnt):
                            for col in range(col_cnt):
                                try:
                                    cell_text = element[row][col]
                                except:
                                    cell_text  = ''
                                table.cell(row, col).text = cell_text

            output_dir = self._get_and_create_relative_output_dir(file_path)
            output_docs_file_path = os.path.join(output_dir,os.path.basename(file_path).replace('.pdf','.docx'))
            document.save(output_docs_file_path)
        except:
            print("Could not convert PDF to DOCX. Skipping  " + file_path)

    def _get_and_create_relative_output_dir(self, input_path):
        # Creates output directories in out put folder relative to input path
        relative_path = Path(input_path).parent.relative_to(self.input_folder)
        output_docs_file_path = Path(self.output_docx_folder_path).joinpath(relative_path)
        output_docs_file_path.mkdir(parents=True, exist_ok=True)
        return str(output_docs_file_path)
    def _convert_doc_to_docx_and_write(self, doc_file_path):
        try:
            output_dir = os.path.join(self.input_folder,self._get_and_create_relative_output_dir(doc_file_path))
            os.system(f'soffice --headless --convert-to docx "{doc_file_path}" --outdir "{output_dir}"')
        except:
            print("Could not convert DOC to DOCX. Skipping  " + doc_file_path)




    def convert_to_docx(self):
        # 1. Convert PDF files to docx
        Parallel(n_jobs=-1)(
            delayed(self._convert_pdf_to_docx_and_write)(path) for path in
            tqdm(glob.glob(self.input_folder + '/**/*.pdf',recursive=True),desc = 'Converting PDF to DOCX'))

        # 2. Convert doc files to docx
        # Parallel(n_jobs=-1)(
        #     delayed(self._convert_doc_to_docx_and_write)(path) for path in
        #     tqdm(glob.glob(self.input_folder + '/**/*.doc',recursive=True),desc = 'Converting DOC to DOCX'))
        for file_path in tqdm(glob.glob(self.input_folder + '/**/*.doc',recursive=True),desc = 'Converting DOC to DOCX'):
            self._convert_doc_to_docx_and_write(file_path)

        # 3. copy docx files to output folder
        for file_path in tqdm(glob.glob(self.input_folder + '/**/*.docx',recursive=True),desc = "Copying DOCX files"):
            output_dir = self._get_and_create_relative_output_dir(file_path)
            os.system(f'cp "{file_path}" "{output_dir}"')




